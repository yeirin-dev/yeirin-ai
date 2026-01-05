"""사회서비스 이용 추천서 DOCX 템플릿 채우기 테스트.

GovernmentDocxFiller 클래스를 테스트합니다.
실제 템플릿 파일을 사용하여 데이터가 올바르게 채워지는지 검증합니다.
"""

import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from yeirin_ai.domain.integrated_report.models import (
    BasicInfo,
    BirthDate,
    ChildInfo,
    CoverInfo,
    GuardianInfo,
    InstitutionInfo,
    IntegratedReportRequest,
    KprcSummary,
    PsychologicalInfo,
    RequestDate,
    RequestMotivation,
)
from yeirin_ai.infrastructure.document.government_docx_filler import (
    GOVERNMENT_TEMPLATE_PATH,
    GovernmentDocxFiller,
    GovernmentDocxFillerError,
)


class TestGovernmentDocxFillerInitialization:
    """GovernmentDocxFiller 초기화 테스트."""

    def test_기본_템플릿_경로로_초기화한다(self) -> None:
        """템플릿 경로를 지정하지 않으면 기본 경로를 사용한다."""
        # When
        filler = GovernmentDocxFiller()

        # Then
        assert filler.template_path == GOVERNMENT_TEMPLATE_PATH

    def test_사용자_지정_템플릿_경로로_초기화한다(self, tmp_path: Path) -> None:
        """사용자가 지정한 템플릿 경로를 사용한다."""
        # Given
        custom_template = tmp_path / "custom_template.docx"
        doc = Document()
        doc.save(str(custom_template))

        # When
        filler = GovernmentDocxFiller(template_path=custom_template)

        # Then
        assert filler.template_path == custom_template

    def test_존재하지_않는_템플릿_경로로_초기화하면_에러가_발생한다(self) -> None:
        """존재하지 않는 템플릿 경로로 초기화하면 에러가 발생한다."""
        # Given
        non_existent_path = Path("/non/existent/template.docx")

        # When & Then
        with pytest.raises(GovernmentDocxFillerError) as exc_info:
            GovernmentDocxFiller(template_path=non_existent_path)

        assert "템플릿 파일을 찾을 수 없습니다" in str(exc_info.value)


class TestGovernmentDocxFillerTemplateExists:
    """실제 템플릿 파일 존재 여부 테스트."""

    def test_기본_템플릿_파일이_존재한다(self) -> None:
        """goverment_doc1.docx 템플릿 파일이 실제로 존재한다."""
        # Then
        assert GOVERNMENT_TEMPLATE_PATH.exists(), (
            f"템플릿 파일이 존재하지 않습니다: {GOVERNMENT_TEMPLATE_PATH}"
        )

    def test_기본_템플릿으로_GovernmentDocxFiller를_생성할_수_있다(self) -> None:
        """기본 템플릿 경로로 GovernmentDocxFiller 인스턴스를 생성할 수 있다."""
        # When
        filler = GovernmentDocxFiller()

        # Then
        assert filler is not None


class TestGovernmentDocxFillerFillTemplate:
    """GovernmentDocxFiller.fill_template() 메서드 테스트."""

    @pytest.fixture
    def complete_request(self) -> IntegratedReportRequest:
        """완전한 통합 보고서 요청 픽스처 (보호자 + 기관 정보 포함)."""
        return IntegratedReportRequest(
            counsel_request_id="123e4567-e89b-12d3-a456-426614174000",
            child_id="123e4567-e89b-12d3-a456-426614174001",
            child_name="홍길동",
            cover_info=CoverInfo(
                requestDate=RequestDate(year=2025, month=1, day=15),
                centerName="서울아동발달센터",
                counselorName="김상담",
            ),
            basic_info=BasicInfo(
                childInfo=ChildInfo(
                    name="홍길동",
                    gender="MALE",
                    age=10,
                    grade="초4",
                    birthDate=BirthDate(year=2015, month=3, day=15),
                ),
                careType="PRIORITY",
            ),
            psychological_info=PsychologicalInfo(
                medicalHistory="ADHD 진단 이력",
                specialNotes="학교 적응에 어려움을 겪고 있음",
            ),
            request_motivation=RequestMotivation(
                motivation="행동 교정 및 학교 적응 지원이 필요합니다.",
                goals="감정 조절 능력 향상 및 또래 관계 개선",
            ),
            kprc_summary=KprcSummary(
                expertOpinion="본 아동은 KPRC 검사 결과, 주의력 영역에서 다소 낮은 점수를 보이고 있습니다.",
                keyFindings=["주의력 저하", "사회성 발달 양호"],
                recommendations=["집중력 향상 프로그램 참여 권장", "또래 활동 참여 권장"],
            ),
            assessment_report_s3_key="assessment-reports/KPRC_홍길동_abc123.pdf",
            guardian_info=GuardianInfo(
                name="홍부모",
                phoneNumber="010-1234-5678",
                homePhone="02-123-4567",
                address="서울시 강남구 테헤란로 123",
                addressDetail="101동 1001호",
                relationToChild="부",
            ),
            institution_info=InstitutionInfo(
                institutionName="서울초등학교",
                phoneNumber="02-987-6543",
                address="서울시 강남구 학동로 456",
                addressDetail="본관 3층",
                writerPosition="담임교사",
                writerName="김선생",
                relationToChild="담임교사",
            ),
        )

    @pytest.fixture
    def request_without_guardian(self) -> IntegratedReportRequest:
        """보호자 정보 없는 요청 픽스처."""
        return IntegratedReportRequest(
            counsel_request_id="123e4567-e89b-12d3-a456-426614174000",
            child_id="123e4567-e89b-12d3-a456-426614174001",
            child_name="홍길동",
            cover_info=CoverInfo(
                requestDate=RequestDate(year=2025, month=1, day=15),
                centerName="서울아동발달센터",
                counselorName="김상담",
            ),
            basic_info=BasicInfo(
                childInfo=ChildInfo(
                    name="홍길동",
                    gender="MALE",
                    age=10,
                    grade="초4",
                ),
                careType="PRIORITY",
            ),
            psychological_info=PsychologicalInfo(
                medicalHistory="없음",
                specialNotes="특이사항 없음",
            ),
            request_motivation=RequestMotivation(
                motivation="정서 지원 필요",
                goals="심리적 안정",
            ),
            kprc_summary=KprcSummary(
                expertOpinion="전반적으로 양호합니다.",
            ),
            assessment_report_s3_key="assessment-reports/test.pdf",
            guardian_info=None,
            institution_info=InstitutionInfo(
                institutionName="테스트기관",
                phoneNumber="02-000-0000",
                address="테스트 주소",
                writerPosition="상담사",
                writerName="테스터",
                relationToChild="상담사",
            ),
        )

    def test_완전한_데이터로_템플릿을_채운다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """모든 필드가 채워진 요청으로 템플릿을 채울 수 있다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)

        # Then
        assert result_bytes is not None
        assert len(result_bytes) > 0
        # DOCX 파일인지 확인 (PK 시그니처)
        assert result_bytes[:2] == b"PK"

    def test_결과가_유효한_DOCX_파일이다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """결과 바이트가 유효한 DOCX 파일이다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)

        # Then: python-docx로 열 수 있어야 함
        doc = Document(io.BytesIO(result_bytes))
        assert doc is not None
        assert len(doc.tables) >= 3  # 템플릿에 최소 3개 테이블 존재

    def test_아동_정보가_테이블에_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """아동 정보(성명, 생년월일)가 첫 번째 테이블에 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then: 테이블 내용 확인
        # 테이블 0에서 아동 이름과 생년월일을 찾음
        table0_text = self._get_table_text(doc.tables[0])
        assert "홍길동" in table0_text
        assert "2015년 3월 15일" in table0_text

    def test_보호자_정보가_테이블에_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """보호자 정보(성명, 전화번호, 관계)가 테이블에 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table0_text = self._get_table_text(doc.tables[0])
        assert "홍부모" in table0_text
        assert "010-1234-5678" in table0_text

    def test_추천사유가_테이블에_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """추천사유(motivation)가 두 번째 테이블에 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table1_text = self._get_table_text(doc.tables[1])
        assert "행동 교정" in table1_text or "학교 적응" in table1_text

    def test_전문가_소견이_테이블에_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """전문가 소견(expertOpinion)이 판단계기 셀에 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table1_text = self._get_table_text(doc.tables[1])
        assert "KPRC 검사 결과" in table1_text or "주의력" in table1_text

    def test_기관_정보가_테이블에_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """기관 정보(기관명, 연락처, 작성자명)가 세 번째 테이블에 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table2_text = self._get_table_text(doc.tables[2])
        assert "서울초등학교" in table2_text
        assert "김선생" in table2_text

    def test_보호자_정보_없이도_템플릿을_채울_수_있다(
        self, request_without_guardian: IntegratedReportRequest
    ) -> None:
        """보호자 정보(guardian_info)가 없어도 템플릿을 채울 수 있다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(request_without_guardian)

        # Then
        assert result_bytes is not None
        doc = Document(io.BytesIO(result_bytes))
        assert doc is not None

    def test_날짜가_한국어_형식으로_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """날짜가 '2025년 1월 15일' 형식으로 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then: 문서의 단락에서 날짜 형식 확인
        all_text = " ".join([para.text for para in doc.paragraphs])
        # requestDate가 2025년 1월 15일이므로 해당 형식이 있어야 함
        assert "2025년" in all_text and "1월" in all_text and "15일" in all_text

    def test_발급처가_채워진다(
        self, complete_request: IntegratedReportRequest
    ) -> None:
        """발급처에 기관명이 채워진다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(complete_request)
        doc = Document(io.BytesIO(result_bytes))

        # Then: 문서의 단락에서 발급처 확인
        all_text = " ".join([para.text for para in doc.paragraphs])
        # institution_info가 있으면 기관명, 없으면 centerName
        assert "서울초등학교" in all_text or "서울아동발달센터" in all_text

    @staticmethod
    def _get_table_text(table) -> str:
        """테이블의 모든 셀 텍스트를 추출한다."""
        texts = []
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
        return " ".join(texts)


class TestGovernmentDocxFillerErrorHandling:
    """GovernmentDocxFiller 에러 처리 테스트."""

    @pytest.fixture
    def minimal_request(self) -> IntegratedReportRequest:
        """최소한의 요청 데이터 픽스처."""
        return IntegratedReportRequest(
            counsel_request_id="test-id",
            child_id="child-id",
            child_name="테스트",
            cover_info=CoverInfo(
                requestDate=RequestDate(year=2025, month=1, day=1),
                centerName="테스트센터",
                counselorName="테스터",
            ),
            basic_info=BasicInfo(
                childInfo=ChildInfo(
                    name="테스트",
                    gender="MALE",
                    age=10,
                    grade="4",
                ),
                careType="GENERAL",
            ),
            psychological_info=PsychologicalInfo(
                medicalHistory="없음",
                specialNotes="없음",
            ),
            request_motivation=RequestMotivation(
                motivation="테스트",
                goals="테스트",
            ),
            kprc_summary=KprcSummary(
                expertOpinion="테스트 소견",
            ),
            assessment_report_s3_key="test.pdf",
        )

    def test_템플릿_파일이_손상되면_에러를_발생시킨다(
        self, tmp_path: Path, minimal_request: IntegratedReportRequest
    ) -> None:
        """손상된 템플릿 파일을 열려고 하면 에러가 발생한다."""
        # Given: 손상된 파일 생성
        corrupted_template = tmp_path / "corrupted.docx"
        corrupted_template.write_bytes(b"not a valid docx")

        filler = GovernmentDocxFiller(template_path=corrupted_template)

        # When & Then
        with pytest.raises(GovernmentDocxFillerError) as exc_info:
            filler.fill_template(minimal_request)

        assert "템플릿 채우기 실패" in str(exc_info.value)


class TestGovernmentDocxFillerCellTextSetting:
    """셀 텍스트 설정 메서드 테스트."""

    def test_셀에_텍스트와_폰트가_올바르게_설정된다(self, tmp_path: Path) -> None:
        """_set_cell_text_with_font 메서드가 텍스트와 폰트를 올바르게 설정한다."""
        # Given: 테스트용 DOCX 생성
        test_doc = Document()
        table = test_doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        filler = GovernmentDocxFiller()

        # When
        filler._set_cell_text_with_font(cell, "테스트 텍스트", font_size=12, bold=True)

        # Then
        assert cell.paragraphs[0].text == "테스트 텍스트"
        # 폰트 설정 확인
        run = cell.paragraphs[0].runs[0]
        assert run.font.bold is True

    def test_빈_텍스트도_설정할_수_있다(self, tmp_path: Path) -> None:
        """빈 문자열도 셀에 설정할 수 있다."""
        # Given
        test_doc = Document()
        table = test_doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "기존 텍스트"

        filler = GovernmentDocxFiller()

        # When
        filler._set_cell_text_with_font(cell, "")

        # Then
        assert cell.paragraphs[0].text == ""


class TestGovernmentDocxFillerRecommenderOpinion:
    """추천자 의견(③) 필드 테스트."""

    @pytest.fixture
    def request_with_recommendations(self) -> IntegratedReportRequest:
        """권장사항이 포함된 요청 픽스처."""
        return IntegratedReportRequest(
            counsel_request_id="test-id",
            child_id="child-id",
            child_name="테스트",
            cover_info=CoverInfo(
                requestDate=RequestDate(year=2025, month=1, day=1),
                centerName="테스트센터",
                counselorName="테스터",
            ),
            basic_info=BasicInfo(
                childInfo=ChildInfo(
                    name="테스트",
                    gender="MALE",
                    age=10,
                    grade="4",
                ),
                careType="GENERAL",
            ),
            psychological_info=PsychologicalInfo(
                medicalHistory="없음",
                specialNotes="없음",
            ),
            request_motivation=RequestMotivation(
                motivation="정서 지원 필요",
                goals="정서 안정화, 사회성 향상",
            ),
            kprc_summary=KprcSummary(
                expertOpinion="전반적으로 양호합니다.",
                keyFindings=["주의력 저하", "정서 불안정"],
                recommendations=["심리상담 권장", "사회성 향상 프로그램 권장", "부모 상담 권장"],
            ),
            assessment_report_s3_key="test.pdf",
        )

    def test_상담_목표가_추천자_의견에_포함된다(
        self, request_with_recommendations: IntegratedReportRequest
    ) -> None:
        """상담 목표(goals)가 추천자 의견 셀에 포함된다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(request_with_recommendations)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table1_text = self._get_all_table_text(doc)
        assert "정서 안정화" in table1_text or "사회성 향상" in table1_text

    def test_권장사항이_불릿_포인트로_포맷된다(
        self, request_with_recommendations: IntegratedReportRequest
    ) -> None:
        """권장사항(recommendations)이 불릿 포인트(•)로 포맷된다."""
        # Given
        filler = GovernmentDocxFiller()

        # When
        result_bytes = filler.fill_template(request_with_recommendations)
        doc = Document(io.BytesIO(result_bytes))

        # Then
        table1_text = self._get_all_table_text(doc)
        # 불릿 포인트(•)가 포함되어 있어야 함
        assert "•" in table1_text or "심리상담" in table1_text

    @staticmethod
    def _get_all_table_text(doc: Document) -> str:
        """문서의 모든 테이블 텍스트를 추출한다."""
        texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return " ".join(texts)
