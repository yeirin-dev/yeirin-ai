"""사회서비스 이용 추천서 DOCX 템플릿 채우기.

goverment_doc1.docx 템플릿을 데이터로 채웁니다.
"""

import io
import logging
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from yeirin_ai.domain.integrated_report.models import IntegratedReportRequest
from yeirin_ai.infrastructure.llm.recommender_opinion_generator import RecommenderOpinion

logger = logging.getLogger(__name__)

# 나눔고딕 폰트 이름
NANUM_GOTHIC_FONT = "나눔고딕"

# 템플릿 파일 경로
GOVERNMENT_TEMPLATE_PATH = Path(__file__).parent.parent.parent / "goverment_doc1.docx"


class GovernmentDocxFillerError(Exception):
    """사회서비스 이용 추천서 채우기 에러."""

    pass


class GovernmentDocxFiller:
    """사회서비스 이용 추천서 DOCX 템플릿 채우기.

    goverment_doc1.docx 템플릿의 테이블 구조:
    - 테이블 0: 대상자 인적사항
      - Row 0: 성명 | 생년월일 | 주소
      - Row 1: (값) | (값) | (값)
      - Row 2: 보호자 성명 | 이용자와의 관계 | 전화번호(자택/휴대전화)
      - Row 3: (값) | (값) | (자택)/(휴대전화)
    - 테이블 1: 추천사유 및 판단근거
      - Row 0: ① 추천사유
      - Row 1: ② 판단계기
      - Row 2-3: ③ 추천자 의견
    - 테이블 2: 작성자
      - Row 0: 소속기관명 | | | 연락처 |
      - Row 1: 기관소재지 | | | |
      - Row 2: 직 또는 자격 | | 성명 | 서명 또는 날인 |
      - Row 3: 이용자와의 관계 | | | |
    """

    def __init__(self, template_path: Path | None = None) -> None:
        """초기화.

        Args:
            template_path: 템플릿 파일 경로. None이면 기본 경로 사용.
        """
        self.template_path = template_path or GOVERNMENT_TEMPLATE_PATH

        if not self.template_path.exists():
            raise GovernmentDocxFillerError(
                f"템플릿 파일을 찾을 수 없습니다: {self.template_path}"
            )

    def fill_template(
        self,
        request: IntegratedReportRequest,
        recommender_opinion: RecommenderOpinion | None = None,
    ) -> bytes:
        """템플릿을 데이터로 채웁니다.

        Args:
            request: 통합 보고서 생성 요청 데이터
            recommender_opinion: AI 생성 추천자 의견 (선택). 제공되면 ③ 추천자 의견에 사용.

        Returns:
            채워진 DOCX 파일의 바이트 데이터

        Raises:
            GovernmentDocxFillerError: 템플릿 채우기 실패 시
        """
        try:
            doc = Document(str(self.template_path))

            # 테이블 0: 대상자 인적사항
            if len(doc.tables) > 0:
                self._fill_personal_info_table(doc.tables[0], request)

            # 테이블 1: 추천사유 및 판단근거
            if len(doc.tables) > 1:
                self._fill_recommendation_reason_table(
                    doc.tables[1], request, recommender_opinion
                )

            # 테이블 2: 작성자
            if len(doc.tables) > 2:
                self._fill_writer_table(doc.tables[2], request)

            # 날짜 채우기 (단락에서)
            self._fill_date(doc, request)

            # 발급처 채우기 (단락에서)
            self._fill_issuer(doc, request)

            # 바이트로 변환
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(
                "사회서비스 이용 추천서 템플릿 채우기 완료",
                extra={
                    "counsel_request_id": request.counsel_request_id,
                    "child_name": request.child_name,
                },
            )

            return buffer.getvalue()

        except Exception as e:
            logger.error(
                "사회서비스 이용 추천서 템플릿 채우기 실패",
                extra={"error": str(e), "counsel_request_id": request.counsel_request_id},
            )
            raise GovernmentDocxFillerError(f"템플릿 채우기 실패: {e}") from e

    def _fill_personal_info_table(self, table, request: IntegratedReportRequest) -> None:
        """테이블 0: 대상자 인적사항을 채웁니다.

        테이블 구조:
        - Row 0: 성명 | 생년월일 | 주소 (헤더)
        - Row 1: (값) | (값) | (값)
        - Row 2: 보호자 성명 | 이용자와의 관계 | 전화번호 (헤더)
        - Row 3: (값) | (값) | (자택)/(휴대전화)
        """
        child = request.basic_info.childInfo
        guardian = request.guardian_info

        # Row 1: 아동 정보
        if len(table.rows) > 1:
            row1 = table.rows[1]
            # 성명
            if len(row1.cells) > 0:
                self._set_cell_text_with_font(row1.cells[0], child.name, font_size=10)
            # 생년월일
            if len(row1.cells) > 1:
                birth_date_str = ""
                if child.birthDate:
                    birth_date_str = child.birthDate.to_korean_string()
                self._set_cell_text_with_font(row1.cells[1], birth_date_str, font_size=10)
            # 주소
            if len(row1.cells) > 2 and guardian:
                address = guardian.address
                if guardian.addressDetail:
                    address = f"{address} {guardian.addressDetail}"
                self._set_cell_text_with_font(row1.cells[2], address, font_size=10)

        # Row 3: 보호자 정보
        if len(table.rows) > 3 and guardian:
            row3 = table.rows[3]
            # 보호자 성명
            if len(row3.cells) > 0:
                self._set_cell_text_with_font(row3.cells[0], guardian.name, font_size=10)
            # 이용자와의 관계
            if len(row3.cells) > 1:
                self._set_cell_text_with_font(
                    row3.cells[1], guardian.relationToChild, font_size=10
                )
            # 전화번호 (자택/휴대전화)
            if len(row3.cells) > 2:
                home_phone = guardian.homePhone or ""
                mobile_phone = guardian.phoneNumber
                phone_str = f"{home_phone}/{mobile_phone}"
                self._set_cell_text_with_font(row3.cells[2], phone_str, font_size=10)

    def _fill_recommendation_reason_table(
        self,
        table,
        request: IntegratedReportRequest,
        recommender_opinion: RecommenderOpinion | None = None,
    ) -> None:
        """테이블 1: 추천사유 및 판단근거를 채웁니다.

        테이블 구조:
        - Row 0: ① 추천사유 | (내용)
        - Row 1: ② 판단계기 | (내용)
        - Row 2-3: ③ 추천자 의견 | (내용)

        Args:
            table: DOCX 테이블 객체
            request: 통합 보고서 요청 데이터
            recommender_opinion: AI 생성 추천자 의견 (선택)
        """
        motivation = request.request_motivation
        psych = request.psychological_info
        kprc = request.kprc_summary

        # Row 0, Col 1: ① 추천사유
        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[0].cells[1], motivation.motivation, font_size=10
            )

        # Row 1, Col 1: ② 판단계기 (관찰내용, 검사결과 등)
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            judgment_basis_parts = []

            # 심리검사 결과
            if kprc.expertOpinion:
                judgment_basis_parts.append(f"[심리검사 결과]\n{kprc.expertOpinion}")

            # 관찰 내용 (특이사항)
            if psych.specialNotes:
                judgment_basis_parts.append(f"\n\n[관찰내용]\n{psych.specialNotes}")

            # 병력
            if psych.medicalHistory and psych.medicalHistory != "없음":
                judgment_basis_parts.append(f"\n\n[기존 병력]\n{psych.medicalHistory}")

            judgment_basis = "".join(judgment_basis_parts) if judgment_basis_parts else ""
            self._set_cell_text_with_font(
                table.rows[1].cells[1], judgment_basis, font_size=10
            )

        # Row 2-3: ③ 추천자 의견 (서비스 지원이 필요한 분야 등)
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            # AI 생성 추천자 의견이 있으면 사용, 없으면 기존 로직 (fallback)
            if recommender_opinion and recommender_opinion.opinion_text:
                # AI 생성 추천자 의견 사용
                opinion = recommender_opinion.opinion_text

                # 필요 서비스 분야가 있으면 추가
                if recommender_opinion.service_needs:
                    needs_text = "\n".join(
                        f"• {need}" for need in recommender_opinion.service_needs
                    )
                    opinion += f"\n\n[서비스 지원이 필요한 분야]\n{needs_text}"

                logger.info(
                    "AI 생성 추천자 의견 사용",
                    extra={"opinion_length": len(opinion)},
                )
            else:
                # Fallback: 기존 로직 사용
                opinion_parts = []

                # 목표
                if motivation.goals:
                    opinion_parts.append(f"[상담 목표]\n{motivation.goals}")

                # 권장사항
                if kprc.recommendations:
                    recommendations_text = "\n".join(
                        f"• {rec}" for rec in kprc.recommendations
                    )
                    opinion_parts.append(f"\n\n[권장사항]\n{recommendations_text}")

                # 핵심 발견사항
                if kprc.keyFindings:
                    findings_text = "\n".join(f"• {f}" for f in kprc.keyFindings)
                    opinion_parts.append(f"\n\n[핵심 발견사항]\n{findings_text}")

                opinion = "".join(opinion_parts) if opinion_parts else ""

                logger.info(
                    "기존 KPRC 기반 추천자 의견 사용 (fallback)",
                    extra={"opinion_length": len(opinion)},
                )

            self._set_cell_text_with_font(table.rows[2].cells[1], opinion, font_size=10)

    def _fill_writer_table(self, table, request: IntegratedReportRequest) -> None:
        """테이블 2: 작성자 정보를 채웁니다.

        테이블 구조:
        - Row 0: 소속기관명 | (값) | | 연락처 | (값)
        - Row 1: 기관소재지 | (값) | | |
        - Row 2: 직 또는 자격 | (값) | 성명 | (값) | 서명
        - Row 3: 이용자와의 관계 : (값) | | | |
        """
        institution = request.institution_info
        cover = request.cover_info

        # 기관 정보가 없으면 cover_info에서 가져옴
        institution_name = (
            institution.institutionName if institution else cover.centerName
        )
        writer_name = institution.writerName if institution else cover.counselorName

        # Row 0: 소속기관명, 연락처
        if len(table.rows) > 0:
            row0 = table.rows[0]
            # 소속기관명 (Cell 1)
            if len(row0.cells) > 1:
                self._set_cell_text_with_font(row0.cells[1], institution_name, font_size=10)
            # 연락처 (Cell 4)
            if len(row0.cells) > 4 and institution:
                self._set_cell_text_with_font(
                    row0.cells[4], institution.phoneNumber, font_size=10
                )

        # Row 1: 기관소재지
        if len(table.rows) > 1 and institution:
            row1 = table.rows[1]
            if len(row1.cells) > 1:
                address = institution.address
                if institution.addressDetail:
                    address = f"{address} {institution.addressDetail}"
                self._set_cell_text_with_font(row1.cells[1], address, font_size=10)

        # Row 2: 직 또는 자격, 성명
        if len(table.rows) > 2:
            row2 = table.rows[2]
            # 직 또는 자격 (Cell 1)
            if len(row2.cells) > 1 and institution:
                self._set_cell_text_with_font(
                    row2.cells[1], institution.writerPosition, font_size=10
                )
            # 성명 (Cell 3)
            if len(row2.cells) > 3:
                self._set_cell_text_with_font(row2.cells[3], writer_name, font_size=10)

        # Row 3: 이용자와의 관계
        if len(table.rows) > 3 and institution:
            row3 = table.rows[3]
            # 기존 텍스트에서 ':' 뒤에 값 추가
            if len(row3.cells) > 0:
                current_text = row3.cells[0].text
                if ":" in current_text:
                    new_text = current_text.split(":")[0] + ": " + institution.relationToChild
                    self._set_cell_text_with_font(row3.cells[0], new_text, font_size=10)

    def _fill_date(self, doc: Document, request: IntegratedReportRequest) -> None:
        """날짜 단락을 채웁니다.

        "20   .   .   ." 형식의 단락을 찾아서 날짜로 변환합니다.
        """
        request_date = request.cover_info.requestDate

        for para in doc.paragraphs:
            if "20   ." in para.text or "20  ." in para.text:
                date_str = f"{request_date.year}년 {request_date.month}월 {request_date.day}일"
                para.text = date_str
                # 가운데 정렬 및 폰트 설정
                for run in para.runs:
                    run.font.name = NANUM_GOTHIC_FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), NANUM_GOTHIC_FONT)
                    run.font.size = Pt(11)
                break

    def _fill_issuer(self, doc: Document, request: IntegratedReportRequest) -> None:
        """발급처 단락을 채웁니다.

        "발  급  처" 다음 줄에 기관명을 추가합니다.
        """
        institution = request.institution_info
        issuer_name = institution.institutionName if institution else request.cover_info.centerName

        for para in doc.paragraphs:
            if "발  급  처" in para.text or "발 급 처" in para.text:
                # 발급처 텍스트 뒤에 기관명 추가
                para.text = f"발  급  처: {issuer_name}"
                for run in para.runs:
                    run.font.name = NANUM_GOTHIC_FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), NANUM_GOTHIC_FONT)
                    run.font.size = Pt(11)
                break

    def _set_cell_text_with_font(
        self, cell, text: str, font_size: int = 10, bold: bool = False
    ) -> None:
        """셀에 텍스트를 설정하고 나눔고딕 폰트를 적용합니다."""
        cell.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = NANUM_GOTHIC_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), NANUM_GOTHIC_FONT)
        run.font.size = Pt(font_size)
        run.font.bold = bold
