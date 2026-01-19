"""DOCX 템플릿 채우기.

counsel_request_format.docx 템플릿을 상담의뢰지 데이터로 채웁니다.

새 문서 포맷 구조:
- 테이블 0: 표지 정보 (의뢰일자, 센터명, 담당자)
- 테이블 1: 섹션 헤더 "1 기본 정보"
- 테이블 2: 기본 정보 (이름, 성별, 연령, 학년)
- 테이블 3: 센터 이용 기준 + 보호대상 아동 기준 (Row 0-2: 돌봄, Row 3-4: 보호대상)
- 테이블 4: 섹션 헤더 "2 정서·심리 관련 정보"
- 테이블 5: 정서심리 정보 (병력, 특이사항)
- 테이블 6: 섹션 헤더 "3 의뢰동기 및 상담목표"
- 테이블 7: 의뢰 정보 (의뢰동기, 목표)
- 테이블 8: 섹션 헤더 "4 검사 결과 및 AI 기반 종합 소견"
- 테이블 9: 검사 결과 + AI 소견 (12행 구조)
  - Row 0: "1. 표준화 검사 결과 요약" 헤더
  - Row 1: KPRC 헤더
  - Row 2: KPRC 내용
  - Row 3: CRTES-R 헤더
  - Row 4: CRTES-R 내용
  - Row 5: SDQ-A 헤더
  - Row 6: SDQ-A 강점/난점 라벨
  - Row 7: SDQ-A 강점/난점 내용
  - Row 8: "2. AI 기반 아동 마음건강 대화 분석 요약" 헤더
  - Row 9: 대화 분석 내용
  - Row 10: "3. 예이린 AI 기반 통합 전문 소견" 헤더
  - Row 11: 통합 소견 내용
- 테이블 10: 섹션 헤더 "5 보호자 동의 여부"
- 단락: 동의 체크박스
"""

import io
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from yeirin_ai.domain.integrated_report.models import IntegratedReportRequest

logger = logging.getLogger(__name__)

# 나눔고딕 폰트 이름
NANUM_GOTHIC_FONT = "나눔고딕"

# 템플릿 파일 경로 (새 포맷)
TEMPLATE_PATH = Path(__file__).parent.parent.parent / "counsel_request_format.docx"


class DocxFillerError(Exception):
    """DOCX 채우기 에러."""

    pass


class CounselRequestDocxFiller:
    """상담의뢰지 DOCX 템플릿 채우기.

    counsel_request_format.docx 템플릿의 테이블 구조:
    - 테이블 0: 표지 정보 (의뢰일자, 센터명, 담당자)
    - 테이블 1: 섹션 헤더 "1 기본 정보"
    - 테이블 2: 기본 정보 (이름, 성별, 연령, 학년)
    - 테이블 3: 센터 이용 기준 + 보호대상 아동 기준
      - Row 0-2: 돌봄 유형 (우선돌봄/일반/돌봄특례)
      - Row 3: 보호대상 아동 유형 (시설/그룹홈)
      - Row 4: 보호대상 사유 (보호자 이탈/학대/질병·가출/지자체)
    - 테이블 4: 섹션 헤더 "2 정서·심리 관련 정보"
    - 테이블 5: 정서심리 정보 (병력, 특이사항)
    - 테이블 6: 섹션 헤더 "3 의뢰동기 및 상담목표"
    - 테이블 7: 의뢰 정보 (의뢰동기, 목표)
    - 테이블 8: 섹션 헤더 "4 검사 결과 및 AI 기반 종합 소견"
    - 테이블 9: 검사 결과 + AI 소견 (12행 구조)
    - 테이블 10: 섹션 헤더 "5 보호자 동의 여부"
    - 단락: 동의 체크박스
    """

    def __init__(self, template_path: Path | None = None) -> None:
        """초기화.

        Args:
            template_path: 템플릿 파일 경로. None이면 기본 경로 사용.
        """
        self.template_path = template_path or TEMPLATE_PATH

        if not self.template_path.exists():
            raise DocxFillerError(f"템플릿 파일을 찾을 수 없습니다: {self.template_path}")

    def fill_template(self, request: IntegratedReportRequest) -> bytes:
        """템플릿을 데이터로 채웁니다.

        Args:
            request: 통합 보고서 생성 요청 데이터

        Returns:
            채워진 DOCX 파일의 바이트 데이터

        Raises:
            DocxFillerError: 템플릿 채우기 실패 시
        """
        try:
            doc = Document(str(self.template_path))

            # 테이블 너비 자동 조정 (PDF 변환 시 여백 잘림 방지)
            self._fix_table_widths(doc)

            # 테이블 0: 표지 정보 (의뢰일자, 센터명, 담당자)
            self._fill_cover_table(doc.tables[0], request)

            # 테이블 2: 기본 정보 (이름, 성별, 연령, 학년)
            self._fill_basic_info_table(doc.tables[2], request)

            # 테이블 3: 센터 이용 기준 (Row 0-2) + 보호대상 아동 기준 (Row 3-4)
            self._fill_care_type_table(doc.tables[3], request)
            self._fill_protected_child_section(doc.tables[3], request)

            # 테이블 5: 정서심리 정보 (병력, 특이사항)
            self._fill_psychological_info_table(doc.tables[5], request)

            # 테이블 7: 의뢰 정보 (의뢰동기, 목표)
            self._fill_motivation_table(doc.tables[7], request)

            # 테이블 9: 검사 결과 및 AI 소견 (새 12행 구조)
            self._fill_assessment_results_table(doc.tables[9], request)

            # 동의 체크박스 처리 (단락에서)
            self._fill_consent(doc)

            # 바이트로 변환
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(
                "DOCX 템플릿 채우기 완료",
                extra={
                    "counsel_request_id": request.counsel_request_id,
                    "child_name": request.child_name,
                },
            )

            return buffer.getvalue()

        except Exception as e:
            logger.error(
                "DOCX 템플릿 채우기 실패",
                extra={"error": str(e), "counsel_request_id": request.counsel_request_id},
            )
            raise DocxFillerError(f"템플릿 채우기 실패: {e}") from e

    def _fill_cover_table(self, table, request: IntegratedReportRequest) -> None:
        """테이블 0: 표지 정보를 채웁니다 (의뢰일자, 센터명, 담당자).

        테이블 구조:
        - Row 0: 의뢰일자 | 2025년 월 일
        - Row 1: 센터명 | (빈칸)
        - Row 2: 담당자 | (서명)
        """
        cover = request.cover_info
        date_str = cover.requestDate.to_korean_string()

        # Row 0, Cell 1: 의뢰일자
        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
            self._set_cell_text_with_font(table.rows[0].cells[1], date_str, font_size=11)

        # Row 1, Cell 1: 센터명
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            self._set_cell_text_with_font(table.rows[1].cells[1], cover.centerName, font_size=11)

        # Row 2, Cell 1: 담당자
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[2].cells[1], f"{cover.counselorName} (서명)", font_size=11
            )

    def _fill_basic_info_table(self, table, request: IntegratedReportRequest) -> None:
        """테이블 2: 기본 정보를 채웁니다.

        테이블 구조:
        - Row 0: 이름 | (값) | 성별 | (값) | 연령 | (값)
        - Row 1: 학년 | (값) | ... (merged cells)
        """
        child = request.basic_info.childInfo

        # Row 0: 이름, 성별, 연령
        row0 = table.rows[0]
        if len(row0.cells) > 1:
            self._set_cell_text_with_font(row0.cells[1], child.name, font_size=10)
        if len(row0.cells) > 3:
            self._set_cell_text_with_font(
                row0.cells[3], self._gender_to_korean(child.gender), font_size=10
            )
        if len(row0.cells) > 5:
            self._set_cell_text_with_font(row0.cells[5], str(child.age), font_size=10)

        # Row 1: 학년
        row1 = table.rows[1]
        if len(row1.cells) > 1:
            self._set_cell_text_with_font(row1.cells[1], child.grade, font_size=10)

    # 우선돌봄 세부 사유 → 한국어 텍스트 매핑
    PRIORITY_REASON_TEXT_MAP: dict[str, str] = {
        "BASIC_LIVELIHOOD": "기초생활보장 수급권자",
        "LOW_INCOME": "차상위계층 가구의 아동",
        "MEDICAL_AID": "의료급여 수급권자",
        "DISABILITY": "장애가구의 아동 또는 장애 아동",
        "MULTICULTURAL": "다문화가족의 아동",
        "SINGLE_PARENT": "한부모가족의 아동",
        "GRANDPARENT": "조손가구의 아동",
        "EDUCATION_SUPPORT": "초·중·고 교육비 지원 대상 아동",
        "MULTI_CHILD": "자녀가 2명 이상인 가구의 아동",
    }

    def _fill_care_type_table(self, table, request: IntegratedReportRequest) -> None:
        """테이블 3: 센터 이용 기준 체크박스를 처리합니다.

        테이블 구조:
        - Row 0: 센터이용기준 | □ 우선돌봄아동 | Cell[2]: 상세사유1 | Cell[3]: 상세사유2
        - Row 1: 센터이용기준 | □ 일반 아동 | 설명 | 설명
        - Row 2: 센터이용기준 | □ 돌봄특례아동 | 설명 | 설명
        """
        care_type = request.basic_info.careType
        priority_reasons = request.basic_info.priorityReasons or []

        # careType에 따라 해당 행의 체크박스를 체크
        # PRIORITY -> Row 0, GENERAL -> Row 1, SPECIAL -> Row 2
        row_index_map = {
            "PRIORITY": 0,
            "GENERAL": 1,
            "SPECIAL": 2,
        }
        target_row_index = row_index_map.get(care_type, 0)

        for i in range(3):
            if i >= len(table.rows):
                continue
            row = table.rows[i]
            if len(row.cells) > 1:
                cell = row.cells[1]
                text = cell.text
                if i == target_row_index:
                    # 체크 표시로 변경
                    cell.text = text.replace("□", "☑")
                else:
                    # 빈 체크박스 유지
                    cell.text = text.replace("☑", "□")

        # 우선돌봄 아동일 경우 상세 카테고리도 체크 (복수 선택 지원)
        if care_type == "PRIORITY" and priority_reasons:
            self._check_priority_reasons(table, priority_reasons)

    def _check_priority_reasons(self, table, priority_reasons: list[str]) -> None:
        """우선돌봄 세부 사유 체크박스를 처리합니다 (복수 선택 지원).

        Row 0의 Cell 2, Cell 3에 상세 카테고리 체크 항목이 있습니다.
        - Cell 2: 기초생활보장, 차상위계층, 의료급여, 장애가구
        - Cell 3: 다문화가족, 한부모가족, 조손가구, 교육비지원, 자녀2명이상
        """
        # 모든 선택된 사유의 한국어 텍스트 수집
        target_texts: set[str] = set()
        for priority_reason in priority_reasons:
            target_text = self.PRIORITY_REASON_TEXT_MAP.get(priority_reason)
            if target_text:
                target_texts.add(target_text)
            else:
                logger.warning(
                    "알 수 없는 priorityReason 값",
                    extra={"priority_reason": priority_reason},
                )

        if not target_texts:
            return

        # Row 0의 Cell 2, Cell 3에서 해당 텍스트들 찾아 체크
        row = table.rows[0]
        for cell_idx in [2, 3]:
            if cell_idx >= len(row.cells):
                continue
            cell = row.cells[cell_idx]
            cell_text = cell.text

            # 셀 내 각 라인을 확인하여 선택된 사유면 체크
            lines = cell_text.split("\n")
            new_lines = []
            for line in lines:
                checked = False
                for target_text in target_texts:
                    if target_text in line:
                        new_lines.append(line.replace("□", "☑"))
                        checked = True
                        logger.debug(
                            "우선돌봄 세부 사유 체크 완료",
                            extra={"target_text": target_text},
                        )
                        break
                if not checked:
                    new_lines.append(line)
            cell.text = "\n".join(new_lines)

    # 보호대상 아동 유형 → 한국어 텍스트 매핑
    PROTECTED_CHILD_TYPE_TEXT_MAP: dict[str, str] = {
        "CHILD_FACILITY": "아동 양육시설",
        "GROUP_HOME": "공동생활가정(그룹홈)",
    }

    # 보호대상 아동 사유 → 한국어 텍스트 매핑
    PROTECTED_CHILD_REASON_TEXT_MAP: dict[str, str] = {
        "GUARDIAN_ABSENCE": "보호자가 없거나 보호자로부터 이탈",
        "ABUSE": "아동을 학대하는 경우",
        "ILLNESS_RUNAWAY": "보호자의 질병, 가출 등",
        "LOCAL_GOVERNMENT": "지방자치단체장이 보호가 필요하다고 인정한 자",
    }

    def _fill_protected_child_section(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """테이블 3 Row 3-4: 보호대상 아동 기준 체크박스를 처리합니다.

        테이블 구조 (새 포맷):
        - Row 3: 보호대상아동 기준 | □ 아동 양육시설 | □ 공동생활가정(그룹홈)
        - Row 4: 보호대상아동 사유 | □ 보호자 이탈... | □ 학대... | □ 질병... | □ 지자체...
        """
        protected_info = request.basic_info.protectedChildInfo
        if not protected_info:
            return

        # Row 3: 보호대상 아동 유형 (CHILD_FACILITY / GROUP_HOME)
        if protected_info.type and len(table.rows) > 3:
            row3 = table.rows[3]
            target_text = self.PROTECTED_CHILD_TYPE_TEXT_MAP.get(protected_info.type)
            if target_text:
                for cell_idx in range(1, len(row3.cells)):
                    cell = row3.cells[cell_idx]
                    if target_text in cell.text:
                        cell.text = cell.text.replace("□", "☑")
                        logger.debug(
                            "보호대상 아동 유형 체크",
                            extra={"type": protected_info.type, "text": target_text},
                        )
                        break

        # Row 4: 보호대상 아동 사유
        if protected_info.reason and len(table.rows) > 4:
            row4 = table.rows[4]
            target_text = self.PROTECTED_CHILD_REASON_TEXT_MAP.get(protected_info.reason)
            if target_text:
                for cell_idx in range(1, len(row4.cells)):
                    cell = row4.cells[cell_idx]
                    cell_text = cell.text
                    # 셀 내 줄바꿈이 있을 수 있으므로 부분 매칭
                    if target_text[:10] in cell_text:
                        lines = cell_text.split("\n")
                        new_lines = []
                        for line in lines:
                            if target_text[:10] in line:
                                new_lines.append(line.replace("□", "☑"))
                            else:
                                new_lines.append(line)
                        cell.text = "\n".join(new_lines)
                        logger.debug(
                            "보호대상 아동 사유 체크",
                            extra={"reason": protected_info.reason, "text": target_text},
                        )
                        break

    def _fill_psychological_info_table(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """테이블 5: 정서심리 정보를 채웁니다.

        테이블 구조:
        - Row 0: 기존 아동 병력 | (값)
        - Row 1: 병력 외 특이사항 | (값)
        """
        psych = request.psychological_info

        # Row 0, Col 1: 기존 아동 병력
        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[0].cells[1], psych.medicalHistory or "없음", font_size=10
            )

        # Row 1, Col 1: 병력 외 특이사항
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[1].cells[1], psych.specialNotes or "없음", font_size=10
            )

    def _fill_motivation_table(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """테이블 7: 의뢰 정보를 채웁니다.

        테이블 구조:
        - Row 0: 의뢰 동기 | (값)
        - Row 1: 보호자 및 의뢰자의 목표 | (값)
        """
        motivation = request.request_motivation

        # Row 0, Col 1: 의뢰 동기
        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[0].cells[1], motivation.motivation, font_size=10
            )

        # Row 1, Col 1: 보호자 및 의뢰자의 목표
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            self._set_cell_text_with_font(
                table.rows[1].cells[1], motivation.goals, font_size=10
            )

    def _fill_assessment_results_table(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """테이블 9: 검사 결과 및 AI 기반 종합 소견을 채웁니다.

        새 문서 포맷 테이블 구조 (12행):
        - Row 0: "1. 표준화 검사 결과 요약" 헤더
        - Row 1: "1) KPRC" 헤더
        - Row 2: KPRC 내용 (3줄 소견)
        - Row 3: "2) CRTES-R" 헤더
        - Row 4: CRTES-R 내용 (3줄 소견)
        - Row 5: "3) SDQ-A" 헤더
        - Row 6: ["강점(사회지향 행동 관련)", "난점(외현화, 내현화)"] 라벨
        - Row 7: SDQ-A 내용 (2셀: 강점 3줄 / 난점 3줄)
        - Row 8: "2. AI 기반 아동 마음건강 대화 분석 요약" 헤더
        - Row 9: Soul-E 대화 분석 내용
        - Row 10: "3. 예이린 AI 기반 통합 전문 소견" 헤더
        - Row 11: 통합 AI 소견 내용
        """
        if len(table.rows) < 12:
            logger.warning("테이블 9 행 수 부족", extra={"row_count": len(table.rows)})
            return

        # 1) KPRC 소견 (Row 2)
        self._fill_kprc_section(table, request)

        # 2) CRTES-R 소견 (Row 4)
        self._fill_crtes_r_section(table, request)

        # 3) SDQ-A 소견 (Row 7: 강점/난점 분리)
        self._fill_sdq_a_section(table, request)

        # 4) Soul-E 대화 분석 (Row 9)
        self._fill_conversation_analysis_section(table, request)

        # 5) 통합 AI 소견 (Row 11)
        self._fill_integrated_opinion_section(table, request)

    # KPRC 척도명 → 한국어 매핑
    KPRC_SCALE_NAME_MAP: dict[str, str] = {
        "ERS": "자아탄력성",
        "ICN": "비일관성",
        "F": "비전형",
        "VDL": "자기보호",
        "PDL": "타인보호",
        "ANX": "불안",
        "DEP": "우울",
        "SOM": "신체화",
        "DLQ": "비행",
        "HPR": "과잉행동",
        "FAM": "가족관계",
        "SOC": "사회관계",
        "PSY": "정신증",
    }

    def _fill_kprc_section(self, table, request: IntegratedReportRequest) -> None:
        """Row 2: KPRC 검사 소견을 채웁니다.

        바우처 기준 충족 시 첫줄에 위험 척도와 T점수를 표시합니다.
        예: "바우처 기준 충족: 자아탄력성(ERS) 28T, 불안(ANX) 72T"
        """
        if len(table.rows) <= 2:
            return

        cell = table.rows[2].cells[0]
        kprc = request.get_kprc_summary_for_doc()

        # KPRC 검사 결과에서 T점수 및 바우처 기준 정보 가져오기
        kprc_assessment = self._get_kprc_assessment(request)
        voucher_line = self._build_kprc_voucher_line(kprc_assessment)

        if kprc and kprc.summaryLines:
            content_lines: list[str] = []

            # 바우처 기준 충족 시 첫줄에 T점수 정보 추가
            if voucher_line:
                content_lines.append(voucher_line)

            # 3줄 요약 표시 (새 포맷)
            for line in kprc.summaryLines[:3]:
                content_lines.append(f"• {line}")

            content = "\n".join(content_lines)
            self._set_cell_text_with_font(cell, content, font_size=10)
        else:
            self._set_cell_text_with_font(cell, "검사 결과가 없습니다.", font_size=10)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _get_kprc_assessment(self, request: IntegratedReportRequest):
        """KPRC 검사 결과를 가져옵니다."""
        if not request.attached_assessments:
            return None

        for assessment in request.attached_assessments:
            if assessment.assessmentType == "KPRC_CO_SG_E":
                return assessment

        return None

    def _build_kprc_voucher_line(self, kprc_assessment) -> str | None:
        """KPRC 바우처 기준 충족 시 첫줄 텍스트를 생성합니다.

        예: "바우처 기준 충족: 자아탄력성(ERS) 28T, 불안(ANX) 72T"
        """
        if not kprc_assessment:
            return None

        voucher_criteria = kprc_assessment.voucherCriteria
        t_scores = kprc_assessment.kprcTScores

        if not voucher_criteria or not voucher_criteria.meets_criteria:
            return None

        if not t_scores or not voucher_criteria.risk_scales:
            return None

        # 위험 척도별 T점수 수집
        risk_score_parts: list[str] = []
        for scale_name in voucher_criteria.risk_scales:
            korean_name = self.KPRC_SCALE_NAME_MAP.get(scale_name, scale_name)
            t_score = self._get_t_score_by_scale_name(t_scores, scale_name)

            if t_score is not None:
                risk_score_parts.append(f"{korean_name}({scale_name}) {t_score}T")

        if not risk_score_parts:
            return None

        return f"바우처 기준 충족: {', '.join(risk_score_parts)}"

    def _get_t_score_by_scale_name(self, t_scores, scale_name: str) -> int | None:
        """척도명으로 T점수를 가져옵니다."""
        score_field_map = {
            "ERS": "ers_t_score",
            "ICN": "icn_t_score",
            "F": "f_t_score",
            "VDL": "vdl_t_score",
            "PDL": "pdl_t_score",
            "ANX": "anx_t_score",
            "DEP": "dep_t_score",
            "SOM": "som_t_score",
            "DLQ": "dlq_t_score",
            "HPR": "hpr_t_score",
            "FAM": "fam_t_score",
            "SOC": "soc_t_score",
            "PSY": "psy_t_score",
        }

        field_name = score_field_map.get(scale_name)
        if not field_name:
            return None

        return getattr(t_scores, field_name, None)

    def _fill_crtes_r_section(self, table, request: IntegratedReportRequest) -> None:
        """Row 4: CRTES-R 검사 소견을 채웁니다."""
        if len(table.rows) <= 4:
            return

        cell = table.rows[4].cells[0]
        crtes_r = self._get_assessment_summary(request, "CRTES_R")

        if crtes_r and crtes_r.summaryLines:
            content = "\n".join(f"• {line}" for line in crtes_r.summaryLines[:3])
            self._set_cell_text_with_font(cell, content, font_size=10)
        else:
            self._set_cell_text_with_font(cell, "검사 결과가 없습니다.", font_size=10)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _fill_sdq_a_section(self, table, request: IntegratedReportRequest) -> None:
        """Row 7: SDQ-A 검사 소견을 채웁니다 (강점/난점 분리).

        SDQ-A는 2개 셀로 분리:
        - Cell 0: 강점 (사회지향 행동 관련) - 3줄
        - Cell 1: 난점 (외현화, 내현화) - 3줄
        """
        if len(table.rows) <= 7:
            return

        row = table.rows[7]
        sdq_a = self._get_assessment_summary(request, "SDQ_A")

        # SDQ-A 요약은 6줄: 앞 3줄은 강점, 뒤 3줄은 난점
        if sdq_a and sdq_a.summaryLines and len(sdq_a.summaryLines) >= 6:
            # 강점 (첫 3줄)
            if len(row.cells) > 0:
                strengths_content = "\n".join(
                    f"• {line}" for line in sdq_a.summaryLines[:3]
                )
                self._set_cell_text_with_font(row.cells[0], strengths_content, font_size=10)
                for para in row.cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 난점 (뒤 3줄)
            if len(row.cells) > 1:
                difficulties_content = "\n".join(
                    f"• {line}" for line in sdq_a.summaryLines[3:6]
                )
                self._set_cell_text_with_font(row.cells[1], difficulties_content, font_size=10)
                for para in row.cells[1].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # 검사 결과 없음
            if len(row.cells) > 0:
                self._set_cell_text_with_font(row.cells[0], "검사 결과가 없습니다.", font_size=10)
            if len(row.cells) > 1:
                self._set_cell_text_with_font(row.cells[1], "검사 결과가 없습니다.", font_size=10)

    def _fill_conversation_analysis_section(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """Row 9: Soul-E AI 대화 분석 요약을 채웁니다."""
        if len(table.rows) <= 9:
            return

        cell = table.rows[9].cells[0]

        # conversationAnalysis 필드가 있으면 사용
        if hasattr(request, "conversationAnalysis") and request.conversationAnalysis:
            analysis = request.conversationAnalysis
            content_parts: list[str] = []

            # 3줄 요약
            if hasattr(analysis, "summaryLines") and analysis.summaryLines:
                for line in analysis.summaryLines[:3]:
                    content_parts.append(f"• {line}")

            # 전문가 분석 (선택)
            if hasattr(analysis, "expertAnalysis") and analysis.expertAnalysis:
                if content_parts:
                    content_parts.append("")
                content_parts.append(analysis.expertAnalysis)

            if content_parts:
                self._set_cell_text_with_font(cell, "\n".join(content_parts), font_size=10)
            else:
                self._set_cell_text_with_font(cell, "대화 분석 결과가 없습니다.", font_size=10)
        else:
            self._set_cell_text_with_font(cell, "AI 상담사와의 대화 기록이 없습니다.", font_size=10)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _fill_integrated_opinion_section(
        self, table, request: IntegratedReportRequest
    ) -> None:
        """Row 11: 예이린 AI 기반 통합 전문 소견을 채웁니다."""
        if len(table.rows) <= 11:
            return

        cell = table.rows[11].cells[0]
        content_parts: list[str] = []

        # 모든 검사의 전문가 소견 통합
        # KPRC
        kprc = request.get_kprc_summary_for_doc()
        if kprc and kprc.expertOpinion:
            content_parts.append(kprc.expertOpinion)

        # CRTES-R
        crtes_r = self._get_assessment_summary(request, "CRTES_R")
        if crtes_r and crtes_r.expertOpinion:
            if content_parts:
                content_parts.append("")
            content_parts.append(crtes_r.expertOpinion)

        # SDQ-A
        sdq_a = self._get_assessment_summary(request, "SDQ_A")
        if sdq_a and sdq_a.expertOpinion:
            if content_parts:
                content_parts.append("")
            content_parts.append(sdq_a.expertOpinion)

        # 대화 분석 전문가 의견 추가
        if hasattr(request, "conversationAnalysis") and request.conversationAnalysis:
            analysis = request.conversationAnalysis
            if hasattr(analysis, "expertAnalysis") and analysis.expertAnalysis:
                if content_parts:
                    content_parts.append("")
                content_parts.append(analysis.expertAnalysis)

        if content_parts:
            self._set_cell_text_with_font(cell, "\n\n".join(content_parts), font_size=10)
        else:
            # 기본 메시지
            default_text = (
                f"{request.child_name} 아동에 대한 종합적인 심리 평가 결과, "
                "전문 상담 서비스를 통한 정서적 지원이 권장됩니다."
            )
            self._set_cell_text_with_font(cell, default_text, font_size=10)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _get_assessment_summary(self, request: IntegratedReportRequest, assessment_type: str):
        """특정 타입의 검사 요약을 가져옵니다."""
        if not request.attached_assessments:
            return None

        for assessment in request.attached_assessments:
            # CRTES_R 또는 SDQ_A 매칭
            if assessment_type in (assessment.assessmentType or ""):
                return assessment.summary

        return None

    def _fill_consent(self, doc: Document) -> None:
        """동의 체크박스를 처리합니다.

        단락 18: □ 동의      □ 미동의
        """
        for para in doc.paragraphs:
            if "□ 동의" in para.text and "□ 미동의" in para.text:
                # 동의에 체크
                para.text = para.text.replace("□ 동의", "☑ 동의", 1)
                break

    def _gender_to_korean(self, gender: str) -> str:
        """성별을 한국어로 변환합니다."""
        return {
            "MALE": "남",
            "FEMALE": "여",
            "M": "남",
            "F": "여",
        }.get(gender.upper(), gender)

    def _set_cell_text_with_font(
        self, cell, text: str, font_size: int = 10, bold: bool = False
    ) -> None:
        """셀에 텍스트를 설정하고 나눔고딕 폰트를 적용합니다."""
        cell.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = NANUM_GOTHIC_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), NANUM_GOTHIC_FONT)
        run.font.size = Pt(font_size)
        run.font.bold = bold

    def _set_paragraph_text_with_font(
        self, paragraph, text: str, font_size: int = 10, bold: bool = False
    ) -> None:
        """단락에 텍스트를 설정하고 나눔고딕 폰트를 적용합니다."""
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = NANUM_GOTHIC_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), NANUM_GOTHIC_FONT)
        run.font.size = Pt(font_size)
        run.font.bold = bold

    def _fix_table_widths(self, doc: Document) -> None:
        """테이블 너비를 페이지 여백 내로 조정합니다.

        LibreOffice PDF 변환 시 테이블 너비가 페이지를 초과하면
        오른쪽 여백이 잘리는 문제가 발생합니다.
        이 메서드는 tblGrid의 gridCol과 셀 너비를 비례 축소합니다.
        """
        # 페이지 사용 가능 너비 계산 (첫 번째 섹션 기준)
        # DXA/Twips 단위 (1 inch = 1440 dxa/twips, 1 mm ≈ 56.7 dxa)
        # Note: section.page_width는 Twips 객체, .twips로 접근
        section = doc.sections[0]
        page_width_dxa = section.page_width.twips
        left_margin_dxa = section.left_margin.twips
        right_margin_dxa = section.right_margin.twips
        # 안전 마진 2mm (약 113 twips) 추가하여 여유 확보
        safety_margin_dxa = 113
        available_width_dxa = page_width_dxa - left_margin_dxa - right_margin_dxa - safety_margin_dxa

        logger.debug(
            "페이지 설정",
            extra={
                "page_width_mm": section.page_width.mm,
                "available_width_mm": available_width_dxa / 56.7,
            },
        )

        for table_idx, table in enumerate(doc.tables):
            tbl = table._tbl

            # tblGrid에서 그리드 컬럼 너비 확인
            tblGrid = tbl.find(qn("w:tblGrid"))
            if tblGrid is None:
                continue

            gridCols = tblGrid.findall(qn("w:gridCol"))
            if not gridCols:
                continue

            # 그리드 컬럼 너비 합계 계산
            total_grid_width = 0
            grid_widths = []
            for gc in gridCols:
                w = gc.get(qn("w:w"))
                width = int(w) if w else 0
                grid_widths.append(width)
                total_grid_width += width

            # 너비가 페이지를 초과하면 비례 축소
            if total_grid_width > available_width_dxa:
                scale_factor = available_width_dxa / total_grid_width
                logger.debug(
                    f"테이블 {table_idx} 너비 조정",
                    extra={
                        "original_width_mm": total_grid_width / 56.7,
                        "available_width_mm": available_width_dxa / 56.7,
                        "scale_factor": f"{scale_factor:.2f}",
                    },
                )

                # 1. gridCol 너비 축소
                for i, gc in enumerate(gridCols):
                    new_width = int(grid_widths[i] * scale_factor)
                    gc.set(qn("w:w"), str(new_width))

                # 2. 모든 행의 셀 너비 축소
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.tcPr
                        if tcPr is not None:
                            tcW = tcPr.find(qn("w:tcW"))
                            if tcW is not None:
                                w_type = tcW.get(qn("w:type"))
                                w_val = tcW.get(qn("w:w"))
                                if w_type == "dxa" and w_val:
                                    new_width = int(int(w_val) * scale_factor)
                                    tcW.set(qn("w:w"), str(new_width))

                # 3. 테이블 가운데 정렬 추가
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement("w:tblPr")
                    tbl.insert(0, tblPr)
                jc = tblPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    tblPr.append(jc)
                jc.set(qn("w:val"), "center")
